#!/usr/bin/env python3
"""
Generate Air Movement 4 Chord Progression as .docx and .pdf
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. Install with: pip install python-docx")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("reportlab not available. Install with: pip install reportlab")

def create_docx():
    """Create .docx document with chord progression"""
    if not DOCX_AVAILABLE:
        print("Cannot create .docx - python-docx not installed")
        return False
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Air Movement 4 - Chord Progression', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Burt Bacharach / Quincy Jones / French Modern Jazz Style')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.italic = True
    subtitle_format.size = Pt(14)
    
    doc.add_paragraph()  # Spacer
    
    # Overview Section
    doc.add_heading('Overview', 1)
    overview = doc.add_paragraph()
    overview.add_run('Overall Structure: ').bold = True
    overview.add_run('Triad Pairs Creating Extended Jazz Harmony')
    
    concept = doc.add_paragraph()
    concept.add_run('Concept: ').bold = True
    concept.add_run('Each measure uses a triad pair - lower triad (Cello) + upper triad (Viola/Violin II) = sophisticated extended chord')
    
    doc.add_paragraph()  # Spacer
    
    # Summary Chord Progression
    doc.add_heading('Summary Chord Progression (90 Measures)', 1)
    
    sections = [
        ('A Section (M1-20)', 'Gmaj7#11 → Am11 → D9 → Bm7b5-D9 → Cmaj9'),
        ('B Section (M21-40)', 'Em11 → Am11 → D9 → Gmaj7#11 → Em11'),
        ('C Section (M41-60)', 'Gmaj7#11 → Am11 → D9 → Bm7b5-D9 → Cmaj9 (with polyrhythms)'),
        ('D Section (M61-80)', 'Em11 → Gmaj7#11 → Am11 → D9 → Gmaj7#11'),
        ("A' Section (M81-90)", 'Gmaj7#11 → Am11 → Gmaj7#11 (final resolution)')
    ]
    
    for section_name, progression in sections:
        p = doc.add_paragraph()
        p.add_run(f'{section_name}: ').bold = True
        p.add_run(progression)
    
    doc.add_paragraph()  # Spacer
    
    # Detailed Sections
    doc.add_heading('Detailed Harmonic Progressions', 1)
    
    # A Section
    doc.add_heading('A Section - Motif 1 Introduction (Measures 1-20)', 2)
    
    a_section = [
        ('M1-4: Gmaj7#11', 'Lydian, air-like', 'Lower: G-B-D (G major)', 'Upper: B-D-F# (B minor)', 'Result: Gmaj7#11 - Lydian color, air-like, ethereal'),
        ('M5-8: Am11', 'Sophisticated minor', 'Lower: A-C-E (A minor)', 'Upper: C-E-G (C major)', 'Result: Am11 - sophisticated, warm'),
        ('M9-12: D9', 'Dominant with color', 'Lower: D-F#-A (D major)', 'Upper: F#-A-C# (F# minor)', 'Result: D9 - dominant function with modal color'),
        ('M13-16: Bm7b5-D9', 'Tritone substitution', 'Lower: B-D-F (B diminished)', 'Upper: D-F#-A (D major)', 'Result: Tritone substitution - sophisticated jazz harmony'),
        ('M17-20: Cmaj9', 'Sophisticated major', 'Lower: C-E-G (C major)', 'Upper: E-G-B (E minor)', 'Result: Cmaj9 - bright, sophisticated')
    ]
    
    for chord, desc, lower, upper, result in a_section:
        p = doc.add_paragraph()
        p.add_run(f'{chord} ').bold = True
        p.add_run(f'({desc})')
        doc.add_paragraph(lower, style='List Bullet')
        doc.add_paragraph(upper, style='List Bullet')
        doc.add_paragraph(result, style='List Bullet')
        doc.add_paragraph()  # Spacer
    
    # B Section
    doc.add_heading('B Section - Motif 2 Introduction (Measures 21-40)', 2)
    
    b_section = [
        ('M21-24: Em11', 'Modal minor, dance-like', 'Lower: E-G-B (E minor)', 'Upper: G-B-D (G major)', 'Result: Em11 - modal, dance-like, rhythmic'),
        ('M25-28: Am11', 'Sequenced', 'Same as M5-8 but with Motif 2', '', ''),
        ('M29-32: D9', 'Developed', 'Same as M9-12 but with Motif 2 development', '', ''),
        ('M33-36: Gmaj7#11', 'Inverted', 'Return to opening harmony with Motif 2 inversion', '', ''),
        ('M37-40: Em11', 'Variation', 'Final Motif 2 variation', '', '')
    ]
    
    for chord, desc, lower, upper, result in b_section:
        p = doc.add_paragraph()
        p.add_run(f'{chord} ').bold = True
        p.add_run(f'({desc})')
        if lower:
            doc.add_paragraph(lower, style='List Bullet')
        if upper:
            doc.add_paragraph(upper, style='List Bullet')
        if result:
            doc.add_paragraph(result, style='List Bullet')
        doc.add_paragraph()  # Spacer
    
    # C Section
    doc.add_heading('C Section - Development with Polyrhythms (Measures 41-60)', 2)
    c_text = doc.add_paragraph('Same chord progression as A Section, but with:')
    c_text.add_run(' M41-44: 3:2 polyrhythm (Violin I triplets vs. others duplets)')
    doc.add_paragraph('M45-48: 4:3 polyrhythm (Violin I 16th notes vs. others triplets)', style='List Bullet')
    doc.add_paragraph('M49-52: 5:4 polyrhythm (Violin I quintuplets vs. others 16th notes)', style='List Bullet')
    doc.add_paragraph()  # Spacer
    
    # D Section
    doc.add_heading('D Section - Integration (Measures 61-80)', 2)
    doc.add_paragraph('M61-64: Em11 (3:2 polyrhythm)', style='List Bullet')
    doc.add_paragraph('M65-68: Gmaj7#11 (Motif combination)', style='List Bullet')
    doc.add_paragraph('M69-72: Am11 (Motif development)', style='List Bullet')
    doc.add_paragraph('M73-76: D9 (4:3 polyrhythm)', style='List Bullet')
    doc.add_paragraph('M77-80: Gmaj7#11 (Motif resolution)', style='List Bullet')
    doc.add_paragraph()  # Spacer
    
    # A' Section
    doc.add_heading("A' Section - Recapitulation (Measures 81-90)", 2)
    doc.add_paragraph('M81-84: Gmaj7#11 (Motif 1 returns)', style='List Bullet')
    doc.add_paragraph('M85-88: Am11 (Final development)', style='List Bullet')
    doc.add_paragraph('M89-90: Gmaj7#11 (Final resolution with harmonics)', style='List Bullet')
    doc.add_paragraph()  # Spacer
    
    # Harmonic Style Reference
    doc.add_heading('Harmonic Style Reference', 1)
    
    doc.add_heading('Burt Bacharach Characteristics', 2)
    bacharach = [
        'Extended chords: Maj7, Min7, Dom9, Maj9, Min9, 11ths, 13ths',
        'Chromatic voice leading',
        'Unexpected modulations',
        'Sophisticated ii-V-I with substitutions',
        'Major 7th chords for elegance',
        'Diminished chords for color'
    ]
    for item in bacharach:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Quincy Jones Characteristics', 2)
    quincy = [
        'Rich chord voicings',
        'Modal interchange (parallel major/minor)',
        'Secondary dominants (V/V, V/vi, V/ii, V/IV)',
        'Smooth voice leading',
        'Extended chords with color tones'
    ]
    for item in quincy:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('French Modern Jazz Characteristics', 2)
    french = [
        'Modal jazz (Lydian, Mixolydian, Dorian)',
        'Quartal/quintal harmony (stacked 4ths/5ths)',
        'Impressionist colors (Debussy/Ravel influence)',
        'Whole-tone scales',
        'Pentatonic scales',
        'Extended chords'
    ]
    for item in french:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()  # Spacer
    
    # Chord Voicing Principles
    doc.add_heading('Chord Voicing Principles', 1)
    
    doc.add_heading('String Quartet Voicing', 2)
    voicing = [
        'Violin I: Melody (top voice)',
        'Violin II: Upper harmony (3rd, 7th, 9th)',
        'Viola: Middle harmony (5th, 7th, 9th, 11th)',
        'Cello: Bass (root, 3rd, 5th) or walking bass'
    ]
    for item in voicing:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Extended Chord Voicings', 2)
    extended = [
        'Maj9: Root (Cello), 3rd (Viola), 5th (Violin II), 7th (Viola), 9th (Violin I)',
        'Min11: Root (Cello), 3rd (Viola), 5th (Violin II), 7th (Viola), 9th (Violin II), 11th (Violin I)',
        'Dom13: Root (Cello), 3rd (Viola), 7th (Violin II), 9th (Violin II), 13th (Violin I)'
    ]
    for item in extended:
        doc.add_paragraph(item, style='List Bullet')
    
    # Save document
    output_path = 'Air-Mov4-Chord-Progression.docx'
    doc.save(output_path)
    print(f"Created {output_path}")
    return True

def create_pdf():
    """Create .pdf document with chord progression"""
    if not REPORTLAB_AVAILABLE:
        print("Cannot create .pdf - reportlab not installed")
        return False
    
    output_path = 'Air-Mov4-Chord-Progression.pdf'
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title
    story.append(Paragraph('Air Movement 4 - Chord Progression', title_style))
    story.append(Paragraph('<i>Burt Bacharach / Quincy Jones / French Modern Jazz Style</i>', 
                          ParagraphStyle('Subtitle', parent=styles['Normal'], alignment=TA_CENTER, fontSize=14)))
    story.append(Spacer(1, 0.2*inch))
    
    # Overview
    story.append(Paragraph('<b>Overview</b>', heading_style))
    story.append(Paragraph('<b>Overall Structure:</b> Triad Pairs Creating Extended Jazz Harmony', styles['Normal']))
    story.append(Paragraph('<b>Concept:</b> Each measure uses a triad pair - lower triad (Cello) + upper triad (Viola/Violin II) = sophisticated extended chord', styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Summary
    story.append(Paragraph('<b>Summary Chord Progression (90 Measures)</b>', heading_style))
    sections = [
        ('A Section (M1-20)', 'Gmaj7#11 → Am11 → D9 → Bm7b5-D9 → Cmaj9'),
        ('B Section (M21-40)', 'Em11 → Am11 → D9 → Gmaj7#11 → Em11'),
        ('C Section (M41-60)', 'Gmaj7#11 → Am11 → D9 → Bm7b5-D9 → Cmaj9 (with polyrhythms)'),
        ('D Section (M61-80)', 'Em11 → Gmaj7#11 → Am11 → D9 → Gmaj7#11'),
        ("A' Section (M81-90)", 'Gmaj7#11 → Am11 → Gmaj7#11 (final resolution)')
    ]
    
    for section_name, progression in sections:
        story.append(Paragraph(f'<b>{section_name}:</b> {progression}', styles['Normal']))
    
    story.append(Spacer(1, 0.2*inch))
    
    # Detailed Sections
    story.append(Paragraph('<b>Detailed Harmonic Progressions</b>', heading_style))
    
    # A Section
    story.append(Paragraph('<b>A Section - Motif 1 Introduction (Measures 1-20)</b>', heading_style))
    a_section = [
        ('M1-4: Gmaj7#11', 'Lydian, air-like', 'Lower: G-B-D (G major)', 'Upper: B-D-F# (B minor)', 'Result: Gmaj7#11 - Lydian color, air-like, ethereal'),
        ('M5-8: Am11', 'Sophisticated minor', 'Lower: A-C-E (A minor)', 'Upper: C-E-G (C major)', 'Result: Am11 - sophisticated, warm'),
        ('M9-12: D9', 'Dominant with color', 'Lower: D-F#-A (D major)', 'Upper: F#-A-C# (F# minor)', 'Result: D9 - dominant function with modal color'),
        ('M13-16: Bm7b5-D9', 'Tritone substitution', 'Lower: B-D-F (B diminished)', 'Upper: D-F#-A (D major)', 'Result: Tritone substitution - sophisticated jazz harmony'),
        ('M17-20: Cmaj9', 'Sophisticated major', 'Lower: C-E-G (C major)', 'Upper: E-G-B (E minor)', 'Result: Cmaj9 - bright, sophisticated')
    ]
    
    for chord, desc, lower, upper, result in a_section:
        story.append(Paragraph(f'<b>{chord}</b> ({desc})', styles['Normal']))
        story.append(Paragraph(f'• {lower}', styles['Normal']))
        story.append(Paragraph(f'• {upper}', styles['Normal']))
        story.append(Paragraph(f'• {result}', styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
    
    # B Section
    story.append(Paragraph('<b>B Section - Motif 2 Introduction (Measures 21-40)</b>', heading_style))
    b_section = [
        ('M21-24: Em11', 'Modal minor, dance-like', 'Lower: E-G-B (E minor)', 'Upper: G-B-D (G major)', 'Result: Em11 - modal, dance-like, rhythmic'),
        ('M25-28: Am11', 'Sequenced', 'Same as M5-8 but with Motif 2', '', ''),
        ('M29-32: D9', 'Developed', 'Same as M9-12 but with Motif 2 development', '', ''),
        ('M33-36: Gmaj7#11', 'Inverted', 'Return to opening harmony with Motif 2 inversion', '', ''),
        ('M37-40: Em11', 'Variation', 'Final Motif 2 variation', '', '')
    ]
    
    for chord, desc, lower, upper, result in b_section:
        story.append(Paragraph(f'<b>{chord}</b> ({desc})', styles['Normal']))
        if lower:
            story.append(Paragraph(f'• {lower}', styles['Normal']))
        if upper:
            story.append(Paragraph(f'• {upper}', styles['Normal']))
        if result:
            story.append(Paragraph(f'• {result}', styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
    
    # C Section
    story.append(Paragraph('<b>C Section - Development with Polyrhythms (Measures 41-60)</b>', heading_style))
    story.append(Paragraph('Same chord progression as A Section, but with:', styles['Normal']))
    story.append(Paragraph('• M41-44: 3:2 polyrhythm (Violin I triplets vs. others duplets)', styles['Normal']))
    story.append(Paragraph('• M45-48: 4:3 polyrhythm (Violin I 16th notes vs. others triplets)', styles['Normal']))
    story.append(Paragraph('• M49-52: 5:4 polyrhythm (Violin I quintuplets vs. others 16th notes)', styles['Normal']))
    story.append(Spacer(1, 0.1*inch))
    
    # D Section
    story.append(Paragraph('<b>D Section - Integration (Measures 61-80)</b>', heading_style))
    story.append(Paragraph('• M61-64: Em11 (3:2 polyrhythm)', styles['Normal']))
    story.append(Paragraph('• M65-68: Gmaj7#11 (Motif combination)', styles['Normal']))
    story.append(Paragraph('• M69-72: Am11 (Motif development)', styles['Normal']))
    story.append(Paragraph('• M73-76: D9 (4:3 polyrhythm)', styles['Normal']))
    story.append(Paragraph('• M77-80: Gmaj7#11 (Motif resolution)', styles['Normal']))
    story.append(Spacer(1, 0.1*inch))
    
    # A' Section
    story.append(Paragraph("<b>A' Section - Recapitulation (Measures 81-90)</b>", heading_style))
    story.append(Paragraph('• M81-84: Gmaj7#11 (Motif 1 returns)', styles['Normal']))
    story.append(Paragraph('• M85-88: Am11 (Final development)', styles['Normal']))
    story.append(Paragraph('• M89-90: Gmaj7#11 (Final resolution with harmonics)', styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Harmonic Style Reference
    story.append(Paragraph('<b>Harmonic Style Reference</b>', heading_style))
    
    story.append(Paragraph('<b>Burt Bacharach Characteristics</b>', styles['Heading3']))
    bacharach = [
        'Extended chords: Maj7, Min7, Dom9, Maj9, Min9, 11ths, 13ths',
        'Chromatic voice leading',
        'Unexpected modulations',
        'Sophisticated ii-V-I with substitutions',
        'Major 7th chords for elegance',
        'Diminished chords for color'
    ]
    for item in bacharach:
        story.append(Paragraph(f'• {item}', styles['Normal']))
    
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph('<b>Quincy Jones Characteristics</b>', styles['Heading3']))
    quincy = [
        'Rich chord voicings',
        'Modal interchange (parallel major/minor)',
        'Secondary dominants (V/V, V/vi, V/ii, V/IV)',
        'Smooth voice leading',
        'Extended chords with color tones'
    ]
    for item in quincy:
        story.append(Paragraph(f'• {item}', styles['Normal']))
    
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph('<b>French Modern Jazz Characteristics</b>', styles['Heading3']))
    french = [
        'Modal jazz (Lydian, Mixolydian, Dorian)',
        'Quartal/quintal harmony (stacked 4ths/5ths)',
        'Impressionist colors (Debussy/Ravel influence)',
        'Whole-tone scales',
        'Pentatonic scales',
        'Extended chords'
    ]
    for item in french:
        story.append(Paragraph(f'• {item}', styles['Normal']))
    
    story.append(Spacer(1, 0.2*inch))
    
    # Chord Voicing Principles
    story.append(Paragraph('<b>Chord Voicing Principles</b>', heading_style))
    
    story.append(Paragraph('<b>String Quartet Voicing</b>', styles['Heading3']))
    voicing = [
        'Violin I: Melody (top voice)',
        'Violin II: Upper harmony (3rd, 7th, 9th)',
        'Viola: Middle harmony (5th, 7th, 9th, 11th)',
        'Cello: Bass (root, 3rd, 5th) or walking bass'
    ]
    for item in voicing:
        story.append(Paragraph(f'• {item}', styles['Normal']))
    
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph('<b>Extended Chord Voicings</b>', styles['Heading3']))
    extended = [
        'Maj9: Root (Cello), 3rd (Viola), 5th (Violin II), 7th (Viola), 9th (Violin I)',
        'Min11: Root (Cello), 3rd (Viola), 5th (Violin II), 7th (Viola), 9th (Violin II), 11th (Violin I)',
        'Dom13: Root (Cello), 3rd (Viola), 7th (Violin II), 9th (Violin II), 13th (Violin I)'
    ]
    for item in extended:
        story.append(Paragraph(f'• {item}', styles['Normal']))
    
    # Build PDF
    doc.build(story)
    print(f"Created {output_path}")
    return True

if __name__ == '__main__':
    print("Generating Air Movement 4 Chord Progression documents...")
    print()
    
    docx_created = create_docx()
    pdf_created = create_pdf()
    
    print()
    if docx_created and pdf_created:
        print("Successfully created both .docx and .pdf files")
    elif docx_created:
        print("Created .docx file (PDF creation skipped - install reportlab)")
    elif pdf_created:
        print("Created .pdf file (DOCX creation skipped - install python-docx)")
    else:
        print("Could not create documents. Please install:")
        print("  pip install python-docx reportlab")

